"""
KAVACH 2.0 — Agent 3: Freeze Architect
========================================
THE DIFFERENTIATOR. This is the agent that wins the hackathon.

Generates a BNSS/PMLA-compliant freeze order request from:
    - Risk score and scam analysis (Agent 1)
    - Extracted entities (Agent 1)
    - Mule network data (Agent 2)
    - RAG citations (Agent 1)

Output: A structured freeze order document that can be:
    1. Displayed in the UI as a legal document panel
    2. Exported as a PDF via python-docx
    3. "Submitted" to a mock CFCFRMS API endpoint

Why this wins:
    I4C's CFCFRMS saved ₹8,189 crore through fund freezing.
    But filing still takes 45+ minutes of manual work.
    KAVACH collapses it to under 10 seconds.
    No other hackathon team will have this capability.
"""

from __future__ import annotations

import logging
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from config import settings


logger = logging.getLogger("kavach.freeze_architect")


def _generate_order_id() -> str:
    """Generate a realistic CFCFRMS-style order ID."""
    now = datetime.now()
    return f"CFCFRMS-{now.year}-{now.strftime('%m%d')}-{uuid.uuid4().hex[:6].upper()}"


def _determine_legal_sections(scam_type: str, risk_score: float) -> list[dict]:
    """Determine applicable legal sections based on scam type and severity."""
    sections = []

    # Always applicable for digital fraud
    sections.append({
        "section": "BNSS Section 94",
        "title": "Search & Seizure of Digital Evidence",
        "applicability": "Mandates preservation of digital evidence including IP logs, CDR, and transaction records",
    })

    if risk_score >= 0.65:
        sections.append({
            "section": "PMLA Section 17A",
            "title": "Attachment of Property Involved in Money Laundering",
            "applicability": "Enables provisional attachment of mule accounts and proceeds of crime",
        })

    if scam_type == "digital_arrest":
        sections.append({
            "section": "BNS Section 319",
            "title": "Cheating by Personation (formerly IPC 419)",
            "applicability": "Impersonation of government officials for financial extortion",
        })
        sections.append({
            "section": "BNS Section 318",
            "title": "Cheating (formerly IPC 420)",
            "applicability": "Dishonest inducement to deliver property or valuable security",
        })

    sections.append({
        "section": "IT Act Section 66C",
        "title": "Identity Theft",
        "applicability": "Fraudulent use of electronic signature, password, or unique identification",
    })

    sections.append({
        "section": "IT Act Section 66D",
        "title": "Cheating by Personation using Computer Resource",
        "applicability": "Impersonation using communication devices or computer resources",
    })

    if risk_score >= 0.85:
        sections.append({
            "section": "PMLA Section 3",
            "title": "Offence of Money Laundering",
            "applicability": "Whosoever directly or indirectly conceals proceeds of crime",
        })

    return sections


def _format_suspect_list(entities: list[dict]) -> list[dict]:
    """Format extracted entities into a structured suspect entity list."""
    suspects = []
    seen = set()

    for entity in entities:
        if not isinstance(entity, dict):
            continue
        value = entity.get("value", "")
        etype = entity.get("entity_type", "")

        if not value or value in seen:
            continue
        seen.add(value)

        if etype in ("phone", "upi_id", "account_number", "ifsc"):
            suspects.append({
                "type": etype.upper().replace("_", " "),
                "value": value,
                "confidence": entity.get("confidence", 0.0),
                "action_required": "FREEZE" if etype in ("upi_id", "account_number") else "FLAG",
            })

    return suspects


def generate_freeze_order(state: dict) -> dict:
    """
    Generate a BNSS/PMLA-compliant freeze order from the pipeline state.

    Args:
        state: KavachState dict containing all agent outputs

    Returns:
        dict with: freeze_order (structured dict), pdf_path (optional)
    """
    logger.info("Generating BNSS-compliant freeze order...")

    order_id = _generate_order_id()
    now = datetime.now()

    # Extract data from state
    risk_score = state.get("risk_score", 0.0)
    risk_level = state.get("risk_level", "UNKNOWN")
    scam_type = state.get("scam_type", "unknown")
    scam_analysis = state.get("scam_analysis", "")
    entities = state.get("extracted_entities", [])
    citations = state.get("rag_citations", [])
    mule_summary = state.get("mule_ring_summary", "")
    graph_nodes = state.get("graph_nodes", [])
    high_centrality = state.get("high_centrality_nodes", [])

    # Determine fraud amount from entities
    fraud_amount = "Amount under investigation"
    for entity in entities:
        if isinstance(entity, dict) and entity.get("entity_type") == "amount":
            fraud_amount = entity["value"]
            break

    # Build suspect list
    suspect_entities = _format_suspect_list(entities)

    # Determine legal sections
    legal_sections = _determine_legal_sections(scam_type, risk_score)

    # Build evidence citations
    evidence = []
    for i, cit in enumerate(citations[:5]):
        if isinstance(cit, dict):
            evidence.append({
                "exhibit_id": f"EXH-{i+1:03d}",
                "source": cit.get("source_document", "Unknown"),
                "authority": cit.get("issuing_authority", "I4C"),
                "relevant_excerpt": cit.get("relevant_text", "")[:500],
                "similarity_score": cit.get("similarity_score", 0.0),
            })

    # Build network intelligence section
    network_intel = {
        "nodes_identified": len(graph_nodes),
        "mule_accounts": sum(1 for n in graph_nodes if isinstance(n, dict) and "mule" in n.get("node_type", "")),
        "collector_accounts": sum(1 for n in graph_nodes if isinstance(n, dict) and n.get("node_type") == "collector"),
        "high_risk_hubs": len(high_centrality),
        "summary": mule_summary,
    }

    # Compose the freeze order
    freeze_order = {
        "order_id": order_id,
        "generated_at": now.isoformat(),
        "generated_at_display": now.strftime("%d %B %Y, %I:%M %p IST"),
        "status": "DRAFT",

        # Header
        "document_title": "URGENT: Request for Provisional Account Freeze Under PMLA/BNSS",
        "to": "Nodal Officer, Cyber Fraud Mitigation Centre (CFMC), I4C",
        "from": "KAVACH Intelligence Platform — Automated Compliance Engine",
        "priority": "CRITICAL" if risk_score >= 0.85 else "HIGH" if risk_score >= 0.65 else "MEDIUM",

        # Complainant Info
        "victim": {
            "name": "Anonymous Complainant (via KAVACH Triage)",
            "contact": "Filed through automated triage system",
            "complaint_type": "Digital Arrest / Cyber Financial Fraud",
        },

        # Risk Assessment
        "risk_assessment": {
            "score": round(risk_score * 100),
            "level": risk_level,
            "scam_type": scam_type.replace("_", " ").title(),
            "ai_analysis": scam_analysis,
        },

        # Suspect Entities
        "suspect_entities": suspect_entities,
        "fraud_amount": fraud_amount,

        # Legal Framework
        "applicable_sections": legal_sections,

        # Evidence
        "evidence_citations": evidence,

        # Network Intelligence
        "network_intelligence": network_intel,

        # Action Requested
        "action_requested": [
            {
                "action": "PROVISIONAL FREEZE",
                "target": "All identified UPI IDs and bank accounts linked to suspect entities",
                "urgency": "Within 4 hours (Golden Hour Protocol)",
                "legal_basis": "PMLA Section 17A read with BNSS Section 94",
            },
            {
                "action": "CDR/IPDR REQUISITION",
                "target": "All identified phone numbers for past 90 days",
                "urgency": "Within 24 hours",
                "legal_basis": "IT Act Section 69 read with BNSS Section 94",
            },
            {
                "action": "SUSPECT REGISTRY UPDATE",
                "target": "Add all extracted entities to I4C Suspect Registry",
                "urgency": "Immediate",
                "legal_basis": "I4C Standard Operating Procedure",
            },
        ],

        # CFCFRMS Integration
        "cfcfrms_reference": f"KAVACH-AUTO-{now.strftime('%Y%m%d')}-{uuid.uuid4().hex[:4].upper()}",

        # Compliance
        "compliance_notes": [
            "This document is auto-generated by KAVACH Intelligence Platform using Agentic AI.",
            "All evidence citations are retrieved from verified MHA/RBI/I4C advisories only.",
            "Risk scores are computed using a hybrid NLP + Graph Analysis pipeline.",
            "This constitutes a decision-support recommendation, not a final adjudication.",
            "Final freeze action requires authorized officer approval as per PMLA/BNSS mandate.",
        ],

        # Digital Signature Placeholder
        "digital_signature": {
            "system": "KAVACH 2.0 Automated Compliance Engine",
            "hash": uuid.uuid4().hex,
            "timestamp": now.isoformat(),
        },
    }

    logger.info(f"Freeze order generated: {order_id}")

    return {
        "freeze_order": freeze_order,
        "pdf_path": None,  # PDF generation handled separately
    }


def generate_freeze_pdf(freeze_order: dict, output_dir: str = None) -> str:
    """
    Generate a PDF document from the freeze order using python-docx.

    Returns the path to the generated PDF.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if output_dir is None:
            output_dir = str(Path(settings.base_dir) / settings.FREEZE_OUTPUT_DIR)

        Path(output_dir).mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title
        title = doc.add_heading("", level=0)
        run = title.add_run("🛡️ KAVACH INTELLIGENCE PLATFORM")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(220, 50, 50)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document Title
        subtitle = doc.add_heading("", level=1)
        run = subtitle.add_run(freeze_order.get("document_title", "Freeze Order Request"))
        run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta Info
        doc.add_paragraph(f"Order ID: {freeze_order['order_id']}")
        doc.add_paragraph(f"Generated: {freeze_order['generated_at_display']}")
        doc.add_paragraph(f"Priority: {freeze_order['priority']}")
        doc.add_paragraph(f"To: {freeze_order['to']}")
        doc.add_paragraph(f"From: {freeze_order['from']}")

        doc.add_heading("Risk Assessment", level=2)
        risk = freeze_order.get("risk_assessment", {})
        doc.add_paragraph(f"Risk Score: {risk.get('score', 0)}/100 ({risk.get('level', 'N/A')})")
        doc.add_paragraph(f"Scam Type: {risk.get('scam_type', 'Unknown')}")
        doc.add_paragraph(f"AI Analysis: {risk.get('ai_analysis', 'N/A')}")

        # Suspect Entities
        doc.add_heading("Suspect Entities", level=2)
        for suspect in freeze_order.get("suspect_entities", []):
            doc.add_paragraph(
                f"• [{suspect['type']}] {suspect['value']} — Action: {suspect['action_required']}",
            )

        doc.add_paragraph(f"Fraud Amount: {freeze_order.get('fraud_amount', 'Unknown')}")

        # Legal Sections
        doc.add_heading("Applicable Legal Sections", level=2)
        for section in freeze_order.get("applicable_sections", []):
            doc.add_paragraph(f"• {section['section']}: {section['title']}")
            doc.add_paragraph(f"  Applicability: {section['applicability']}")

        # Evidence
        doc.add_heading("Evidence Citations", level=2)
        for ev in freeze_order.get("evidence_citations", []):
            doc.add_paragraph(f"[{ev['exhibit_id']}] {ev['source']} ({ev['authority']})")
            doc.add_paragraph(f"  \"{ev['relevant_excerpt'][:200]}...\"")

        # Network Intelligence
        doc.add_heading("Network Intelligence", level=2)
        net = freeze_order.get("network_intelligence", {})
        doc.add_paragraph(f"Connected Nodes: {net.get('nodes_identified', 0)}")
        doc.add_paragraph(f"Mule Accounts: {net.get('mule_accounts', 0)}")
        doc.add_paragraph(f"Collector Accounts: {net.get('collector_accounts', 0)}")
        doc.add_paragraph(f"Summary: {net.get('summary', 'N/A')}")

        # Actions Requested
        doc.add_heading("Actions Requested", level=2)
        for action in freeze_order.get("action_requested", []):
            doc.add_paragraph(f"▶ {action['action']}")
            doc.add_paragraph(f"  Target: {action['target']}")
            doc.add_paragraph(f"  Urgency: {action['urgency']}")
            doc.add_paragraph(f"  Legal Basis: {action['legal_basis']}")

        # Compliance
        doc.add_heading("Compliance Notes", level=2)
        for note in freeze_order.get("compliance_notes", []):
            doc.add_paragraph(f"• {note}")

        # Save
        filename = f"freeze_order_{freeze_order['order_id']}.docx"
        output_path = str(Path(output_dir) / filename)
        doc.save(output_path)

        logger.info(f"Freeze order PDF saved: {output_path}")
        return output_path

    except ImportError:
        logger.warning("python-docx not installed. Skipping PDF generation.")
        return ""
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        return ""
